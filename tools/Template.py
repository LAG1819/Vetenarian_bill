import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import *
from random import randint
from copy import deepcopy
import random
import numpy as np
import decimal as dc

def make_word(inputAnrede, inputNachname,inputVorname,inputStraße_hausnr,
              inputAddresse, inputEmail, inputBefundListe, inputLeistungen, inputBillCounter,inputPetCounter, inputTag):
    #C:\Users\Pepi\AppData\Local\Rechnung_template\
    #pathR = r"C:\Users\Pepi\AppData\Local\Rechnung_template\RechnungVorlage.docx"
    pathR =r"RechnungVorlage.docx"
    doc = Document(pathR)
    befundStelle = 0
    inputTiername = inputLeistungen[0][6]
    inputTierart = inputLeistungen[0][5]

    if not inputTiername and not inputTiername:
        inputTiername ="Name"
        inputTierart = "No"

    beträgeMedikamente = 0
    beträgeBehandlungen = 0
    beträgeNahrungsergänzung = 0
    newOne= 0
    summe = 0
    mwst_sum=0

    paragraphs = doc.paragraphs
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    #Customerinfo
    # paragraph 1: Anrede
    # paragraph 2: Vorname Nachname
    # paragraph 3: Straße Nr
    # paragraph 5: zip code city
    # paragraph 6: email

    paragraphs[1].text = inputAnrede
    paragraphs[2].text = inputVorname+" "+ inputNachname
    paragraphs[3].text = inputStraße_hausnr
    paragraphs[5].text = inputAddresse
    #paragraphs[0].add_run("Text").italic = True -> kursiv
    #paragraphs[6].text = inputEmail
    paragraphs[6].add_run(inputEmail).italic = True

    #Date and bill no:
    # paragraph 10: Rechnung Nr. 2020/08/31 vom 27.08.2020
    date = datetime.now()
    timestamp = date.strftime("%d.%m.%Y")#(%H:%M:%S.%f) hour, minute, second

    randomNbr = str(randint(0,9))+str(randint(0,9))+str(randint(0,9))
    RNR = inputTierart[0].upper()+inputTiername[0].upper()+"{0:0>2}".format(str(inputBillCounter))+"-"+randomNbr
    paragraphs[10].text = "Rechnung Nr. "+ RNR +" vom "+timestamp

    # paragraph 12: Hund Flora, blutiger Durchfall, Abgabe Probiotikum
    if len(inputBefundListe)>1:
        befundStelle +=1
    try:
        paragraphs[12].add_run(inputTierart + " "+inputTiername + ", "+inputBefundListe[0]).bold = True
    except:
        inputBefundListe[0] =""
        paragraphs[12].add_run(inputTierart + " "+inputTiername + ", "+inputBefundListe[0]).bold = True

    for i in [1,2,3,5,6,10,12]:
        paragraphs[i].style = doc.styles['Normal']
    # table 1: Briefkopf
    # table2: Leistungen (row,col)
    #table 3: Summe Geld
    tableLeistungen = doc.tables[1]

    #cell(1,1) Tierart Name
    tableLeistungen.cell(1,1).paragraphs[0].add_run(inputTierart+" "+inputTiername).bold = True

    #add new row if more then one row is written
    addit_rows = max(0,(len(inputLeistungen)-2)) + inputPetCounter-1
    if(len(inputLeistungen))>1:
        for l in range(addit_rows):
            row = tableLeistungen.rows[2]
            tbl = tableLeistungen._tbl
            border_copied = deepcopy(row._tr)
            tr = border_copied
            row._tr.addnext(tr)

    #add Leistungen to table
    j = 0
    i = 2
    counter = 1
    table_size=len(inputLeistungen)+2
    while i<table_size:
        #check if more than one pet
        if inputTiername != inputLeistungen[j][6]:
            counter +=1
            #new pet
            inputTiername = inputLeistungen[j][6]
            inputTierart = inputLeistungen[j][5]

            tableLeistungen.cell(i,1).paragraphs[0].add_run(inputTierart+" "+inputTiername).bold = True

            #add new Befund Zeile
            new_p = OxmlElement("w:p")
            paragraphs[12]._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraphs[12]._parent)
            new_para.add_run(inputTierart + " "+inputTiername + ", "+inputBefundListe[befundStelle]).bold = True
            try:
                befundStelle+=1
            except:
                pass

            table_size +=1
            i += 1
        #cell(1,0-xxx) Datum der Behandlung

        try:
            date=str(inputLeistungen[j][4])
            year=date.split("/")[2]
            month=date.split("/")[1]
            day=date.split("/")[0]
            eu_date =day+"."+month+"."+year
            tableLeistungen.cell(i,0).text = eu_date
        except:
            tableLeistungen.cell(1,0).text = ""

        #cell(1,2-xxx) Behandlungen
        tableLeistungen.cell(i,1).text = inputLeistungen[j][0]

        #cell(2,2-xxx) Paragraph
        if inputLeistungen[j][8]:
            tableLeistungen.cell(i,2).text = "§ " + inputLeistungen[j][8]

        #cell(4,2-xxx) Preis der Medikamente, prüft ob leer
        if not inputLeistungen[j][1]:
            tableLeistungen.cell(i,3).text = ""
        else:
            if inputLeistungen[j][3] == True:
                beträgeNahrungsergänzung += inputLeistungen[j][1]
            else:
                beträgeMedikamente += inputLeistungen[j][1]

            tableLeistungen.cell(i,3).text = convert_money(np.round(inputLeistungen[j][1],decimals = 2))

        #cell(4,2-xxx) Preis der Behandlungen, prüft ob leer
        if not inputLeistungen[j][2]:
            tableLeistungen.cell(i,4).text =""
        else:
            tableLeistungen.cell(i,4).text = convert_money(np.round(inputLeistungen[j][2],decimals = 2))
            beträgeBehandlungen += inputLeistungen[j][2]

        tableLeistungen.cell(i,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tableLeistungen.cell(i,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        j = j+1
        i = i+1



    # 7% Mwst für Nahrungsergänzung
    mwst7Nahr = np.round(beträgeNahrungsergänzung*0.07,decimals =2)
    mwst19Medi = np.round(beträgeMedikamente*0.19,decimals =2)
    mwst19Beh = np.round(beträgeBehandlungen*0.19,decimals =2)
    summe = np.round(beträgeBehandlungen +beträgeMedikamente+beträgeNahrungsergänzung+mwst7Nahr+mwst19Medi+mwst19Beh,2)


    tableSumme = doc.tables[1]
    if len(inputLeistungen) == 1:
        i = i+1
    #cell(1, 4 letzten) Summe der Einzelpositionen incl MWSt in € | Rechnungsbetrag ges. incl. MWSt in € | darin enthalten MWSt 16% |darin enthalten MWSt 5%
    #tableLeistungen.cell(len(inputLeistungen)+3,1).text = "Summe der Einzelpositionen exkl. MwSt in € "
    sumEinzelposMedi = convert_money(np.round(beträgeMedikamente + beträgeNahrungsergänzung,decimals =2))
    tableSumme.cell(i,1).text = "Summe der Einzelpositionen exkl. MwSt in €"
    tableSumme.cell(i,3).text = sumEinzelposMedi
    tableLeistungen.cell(i,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sumEinzelposBeh = convert_money(np.round(beträgeBehandlungen,decimals =2))
    tableSumme.cell(i,4).text = sumEinzelposBeh
    tableLeistungen.cell(i,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #cell 4
    tableSumme.cell(i+1,1).text = "zzgl. MwSt 19%"
    mwst19Medi = convert_money(np.round(mwst19Medi,decimals =2))
    tableSumme.cell(i+1,3).text = mwst19Medi
    tableLeistungen.cell(i+1,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mwst19Beh = convert_money(np.round(mwst19Beh,decimals =2))
    tableSumme.cell(i+1,4).text = mwst19Beh
    tableLeistungen.cell(i+1,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(i+2,1).text = "zzgl. MwSt 7%"
    if beträgeNahrungsergänzung != 0:
        mwst7Nahr = convert_money(np.round(mwst7Nahr,2))
        tableSumme.cell(i+2,3).text = mwst7Nahr
        tableLeistungen.cell(i+2,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #final sum bold
    tableSumme.cell(i+3,1).paragraphs[0].add_run("Rechnungsbetrag ges. incl. MwSt in €").bold = True
    summe = convert_money(np.round(summe,decimals = 2))
    tableSumme.cell(i+3,4).paragraphs[0].add_run(summe).bold = True
    tableLeistungen.cell(i+3,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tableSumme.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(0,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(1,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(1,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(2,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(2,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tableSumme.cell(3,4).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    name = inputVorname + inputNachname
    return doc,name,timestamp,RNR

def convert_money(inputNr):
    inputNr = str(inputNr)
    inputNr = inputNr.replace(".",",")
    euro = inputNr.split(",")[0]
    try:
        cents = inputNr.split(",")[1]
    except:
        cents ="00"
    if len(cents)<2:
        cents=cents+"0"
    output = euro+","+cents+ " €"
    return output

def add_daily_fee(self,tierart,tiername):
    m = 0
    n = False
    date = None
    l=""
    b=""
    if self.tag.get() == 4:
        l = "Notdienst"
        b = 50.0
        tagParagraph = "3a"
    elif self.tag.get() == 3:
        l = "Feiertag"
        b = 10.0
        tagParagraph = "3a"
    elif self.tag.get() == 2:
        l = "Wochenende"
        b = 15.0
        tagParagraph = "3a"
